import docx

def generate_document(content_list, file_path):
    '''Generate a document from content and save it to a specific path'''
    organised_key_list = ['roles_and_responsibilities', 'required_stuff', 'preferred_stuff']
    document = docx.Document()

    # Setting up Landscape mode for a page
    section = document.sections[0]
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    new_height, new_width = docx.shared.Inches(11), docx.shared.Inches(15)
    section.page_height, section.page_width = new_height, new_width

    capitalised_titles = [title.replace('_', ' ').title()
                          for title in organised_key_list]
    for content in content_list:
        document.add_heading(f"{content['title']} -- {content['company']}", level=4)
        document.add_paragraph(content['introduction'])
        document.add_paragraph(f"Link to the job --> {content['passed_url']}")

        table = document.add_table(rows=2, cols=len(organised_key_list))
        for header_name, header_cell in zip(capitalised_titles, table.rows[0].cells):
            header_cell.text = header_name

        ordered_items = [content[key] for key in organised_key_list]
        for item, cell in zip(ordered_items, table.rows[1].cells):
            cell.text = item

        document.add_page_break()

    document.save(file_path)

